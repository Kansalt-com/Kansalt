"""
Generate ATS-friendly tailored resumes and cover letters.
"""
import re
from io import BytesIO
from datetime import datetime
from typing import List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from utils.logger import get_logger

logger = get_logger(__name__)


class DocumentBuilder:
    """Build tailored DOCX documents (resume, cover letter)."""

    @staticmethod
    def _extract_keywords(job: dict, limit: int = 20) -> List[str]:
        """Extract ATS keywords from job posting."""
        try:
            blob = " ".join([
                job.get("title", ""),
                job.get("description_text", ""),
                " ".join(job.get("tags", []) or [])
            ]).lower()

            words = re.findall(r"[a-z0-9\+#/.-]{3,}", blob)

            stopwords = {
                "and", "the", "with", "for", "to", "in", "of", "a", "an", "on", "or", "as",
                "is", "are", "this", "that", "you", "we", "our", "your", "will", "role",
                "team", "work", "remote", "job", "position", "opportunity", "about",
                "from", "must", "should", "required", "prefer", "experience", "years"
            }

            keywords = []
            seen = set()
            for w in words:
                if w not in stopwords and w not in seen:
                    seen.add(w)
                    keywords.append(w)
                    if len(keywords) >= limit:
                        break

            logger.debug(f"Extracted {len(keywords)} keywords from job")
            return keywords
        except Exception as e:
            logger.error(f"Keyword extraction error: {e}")
            return []

    @staticmethod
    def build_tailored_resume(
        base_resume_text: str,
        candidate_name: str,
        email: str,
        phone: str,
        job: dict,
        selected_skills: Optional[List[str]] = None,
        manual_skills: Optional[List[str]] = None,
    ) -> bytes:
        """
        Build ATS-friendly DOCX resume tailored to a job.
        
        Rules:
        - NO "Target Role", "Source", "Company" lines
        - NO URL garbage
        - Rewrite Professional Summary with job keywords + relevant skills
        - Reorder skills to emphasize relevant ones
        - Add "Key Highlights" section (optional)
        """
        try:
            doc = Document()

            # Base style: ATS-friendly (Calibri, 11pt)
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            # Header
            header = doc.add_paragraph()
            header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = header.add_run(candidate_name or "Candidate")
            run.bold = True
            run.font.size = Pt(14)

            # Contact info
            contact_parts = [x for x in [email, phone] if x]
            if contact_parts:
                contact = " | ".join(contact_parts)
                contact_p = doc.add_paragraph(contact)
                contact_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            doc.add_paragraph("")

            # Professional Summary
            doc.add_paragraph("PROFESSIONAL SUMMARY").runs[0].bold = True

            user_skills = list(selected_skills or [])
            for s in (manual_skills or []):
                if s not in user_skills:
                    user_skills.append(s)

            job_title = job.get("title", "").strip()
            company = job.get("company", "").strip()

            # Build summary with job context
            summary_parts = [
                f"Dedicated professional seeking {job_title or 'a challenging role'}"
            ]
            if company:
                summary_parts[0] += f" at {company}"
            summary_parts[0] += "."

            if user_skills:
                summary_parts.append(f"Core competencies: {', '.join(user_skills[:8])}.")

            summary = " ".join(summary_parts)
            doc.add_paragraph(summary)

            # Key Skills
            doc.add_paragraph("KEY SKILLS").runs[0].bold = True
            all_skills = user_skills + DocumentBuilder._extract_keywords(job, limit=12)
            for skill in all_skills[:15]:
                doc.add_paragraph(skill, style="List Bullet")

            # Professional Experience (from base resume, cleaned)
            doc.add_paragraph("PROFESSIONAL EXPERIENCE").runs[0].bold = True

            skip_keywords = {
                "professional summary", "core skills", "resume", "cv",
                "source:", "company:", "target role:", "generated",
                "http://", "https://", "www."
            }

            exp_lines = []
            for line in base_resume_text.split("\n"):
                clean = line.strip()
                if not clean:
                    continue
                # Skip metadata lines
                if any(skip in clean.lower() for skip in skip_keywords):
                    continue
                # Skip if looks like URL
                if "://" in clean or clean.startswith("http"):
                    continue
                exp_lines.append(clean)

            for line in exp_lines[:30]:  # Limit to 30 lines
                doc.add_paragraph(line)

            # Education section (placeholder)
            doc.add_paragraph("EDUCATION").runs[0].bold = True
            doc.add_paragraph("Degree and institution details (as per original resume)")

            # Certifications (placeholder)
            doc.add_paragraph("CERTIFICATIONS").runs[0].bold = True
            doc.add_paragraph("Relevant certifications (as per original resume)")

            # Footer
            doc.add_paragraph("")
            footer = doc.add_paragraph(
                f"Resume tailored on {datetime.now().strftime('%Y-%m-%d')} | ATS-optimized"
            )
            footer.runs[0].italic = True
            footer.runs[0].font.size = Pt(9)

            bio = BytesIO()
            doc.save(bio)
            logger.info(f"Resume generated for job: {job.get('job_code')}")
            return bio.getvalue()

        except Exception as e:
            logger.error(f"Resume generation error: {e}")
            raise

    @staticmethod
    def build_cover_letter(
        candidate_name: str,
        email: str,
        phone: str,
        job: dict,
        base_resume_text: str,
        selected_skills: Optional[List[str]] = None,
        manual_skills: Optional[List[str]] = None,
    ) -> bytes:
        """
        Build ATS-friendly cover letter (3-5 paragraphs).
        
        Structure:
        1. Opening: role + company
        2. Relevant experience/skills
        3. Why interested + value proposition
        4. Closing + call to action
        """
        try:
            doc = Document()

            # Base style
            style = doc.styles["Normal"]
            style.font.name = "Calibri"
            style.font.size = Pt(11)

            # Header
            header = doc.add_paragraph()
            header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = header.add_run(candidate_name or "Candidate")
            run.bold = True
            run.font.size = Pt(12)

            contact = " | ".join([x for x in [email, phone] if x])
            if contact:
                doc.add_paragraph(contact)

            # Date
            doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))

            doc.add_paragraph("")

            # Salutation
            company = job.get("company", "Hiring Team").strip()
            job_title = job.get("title", "position").strip()
            doc.add_paragraph(f"Dear Hiring Manager at {company},")

            # Paragraph 1: Opening
            p1 = (
                f"I am writing to express my strong interest in the {job_title} position. "
                f"With my background and enthusiasm for this opportunity, I am confident I would be a valuable addition to your team."
            )
            doc.add_paragraph(p1)

            # Paragraph 2: Relevant skills
            user_skills = list(selected_skills or [])
            for s in (manual_skills or []):
                if s not in user_skills:
                    user_skills.append(s)

            if user_skills:
                skills_str = ", ".join(user_skills[:6])
                p2 = (
                    f"My professional experience demonstrates expertise in {skills_str}. "
                    f"I have consistently delivered results and remain committed to continuous learning and excellence."
                )
            else:
                p2 = (
                    "Throughout my career, I have developed a diverse skill set and maintained a commitment to excellence. "
                    "I am eager to apply my experience to make a meaningful contribution to your organization."
                )
            doc.add_paragraph(p2)

            # Paragraph 3: Why interested
            p3 = (
                f"I am particularly excited about the opportunity at {company} because of your commitment to innovation and impact. "
                "I believe my values align with your organizational goals, and I am confident I can contribute to your team's success."
            )
            doc.add_paragraph(p3)

            # Paragraph 4: Call to action
            p4 = (
                "Thank you for considering my application. I would welcome the opportunity to discuss how I can contribute to your team. "
                "Please feel free to reach out at your convenience."
            )
            doc.add_paragraph(p4)

            # Closing
            doc.add_paragraph("Best regards,")
            doc.add_paragraph("")
            doc.add_paragraph("")
            doc.add_paragraph(candidate_name or "Candidate")
            if email:
                doc.add_paragraph(email)

            bio = BytesIO()
            doc.save(bio)
            logger.info(f"Cover letter generated for job: {job.get('job_code')}")
            return bio.getvalue()

        except Exception as e:
            logger.error(f"Cover letter generation error: {e}")
            raise
